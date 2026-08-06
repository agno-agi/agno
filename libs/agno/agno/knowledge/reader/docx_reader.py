import asyncio
from pathlib import Path
from typing import IO, Any, Callable, Iterator, List, Optional, Tuple, Union
from uuid import uuid4

from agno.knowledge.chunking.document import DocumentChunking
from agno.knowledge.chunking.strategy import ChunkingStrategy, ChunkingStrategyType
from agno.knowledge.document.base import Document
from agno.knowledge.image import (
    DEFAULT_IMAGE_BASE_URL,
    save_image_markdown,
)
from agno.knowledge.reader.base import Reader
from agno.knowledge.types import ContentType
from agno.utils.log import log_debug, log_error, log_warning

try:
    from docx import Document as DocxDocument  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
except ImportError:
    raise ImportError("The `python-docx` package is not installed. Please install it via `pip install python-docx`.")

_REL_EMBED = qn("r:embed")
_REL_LINK = qn("r:link")
_REL_ID = qn("r:id")
_W_P = qn("w:p")
_W_TBL = qn("w:tbl")
_W_T = qn("w:t")
_W_TAB = qn("w:tab")
_W_BR = qn("w:br")
_W_CR = qn("w:cr")
_W_DRAWING = qn("w:drawing")
_W_PICT = qn("w:pict")


class DocxReader(Reader):
    """Reader for Doc/Docx files.

    Set ``preserve_images=True`` to extract embedded images, persist them via
    the global image store, and insert inline image links (``![](...)``) at
    their original positions. This is not a full markdown export of the
    document—surrounding text stays plain paragraph text.

    Readers may be reused as singletons. Pass ``content_id`` to each ``read`` /
    ``async_read`` call when preserving images (Knowledge does this automatically).
    """

    def __init__(
        self,
        chunking_strategy: Optional[ChunkingStrategy] = None,
        preserve_images: bool = False,
        image_base_url: str = DEFAULT_IMAGE_BASE_URL,
        **kwargs,
    ):
        if chunking_strategy is None:
            chunk_size = kwargs.get("chunk_size", 5000)
            chunking_strategy = DocumentChunking(chunk_size=chunk_size)
        super().__init__(chunking_strategy=chunking_strategy, **kwargs)
        self.preserve_images = preserve_images
        self.image_base_url = image_base_url

    @classmethod
    def get_supported_chunking_strategies(cls) -> List[ChunkingStrategyType]:
        """Get the list of supported chunking strategies for DOCX readers."""
        return [
            ChunkingStrategyType.DOCUMENT_CHUNKER,
            ChunkingStrategyType.CODE_CHUNKER,
            ChunkingStrategyType.FIXED_SIZE_CHUNKER,
            ChunkingStrategyType.SEMANTIC_CHUNKER,
            ChunkingStrategyType.AGENTIC_CHUNKER,
            ChunkingStrategyType.RECURSIVE_CHUNKER,
        ]

    @classmethod
    def get_supported_content_types(cls) -> List[ContentType]:
        return [ContentType.DOCX, ContentType.DOC]

    def read(
        self,
        file: Union[Path, IO[Any]],
        name: Optional[str] = None,
        content_id: Optional[str] = None,
    ) -> List[Document]:
        """Read a docx file and return a list of documents.

        Args:
            file: Path or file-like object to read.
            name: Optional document name.
            content_id: Knowledge content id for image storage paths. Required when
                ``preserve_images=True`` (passed by Knowledge on each insert).
        """
        try:
            if isinstance(file, Path):
                if not file.exists():
                    raise FileNotFoundError(f"Could not find file: {file}")
                log_debug(f"Reading: {file}")
                docx_document = DocxDocument(str(file))
                doc_name = name or file.stem
            else:
                log_debug(f"Reading uploaded file: {getattr(file, 'name', 'BytesIO')}")
                docx_document = DocxDocument(file)
                doc_name = name or getattr(file, "name", "docx_file").split(".")[0]

            if self.preserve_images:
                doc_content = self._extract_with_images(docx_document, content_id=content_id)
            else:
                doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

            documents = [
                Document(
                    name=doc_name,
                    id=str(uuid4()),
                    content=doc_content,
                )
            ]
            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents
            return documents

        except Exception as e:
            log_error(f"Error reading file: {str(e)}")
            return []

    async def async_read(
        self,
        file: Union[Path, IO[Any]],
        name: Optional[str] = None,
        content_id: Optional[str] = None,
    ) -> List[Document]:
        """Asynchronously read a docx file and return a list of documents"""
        try:
            return await asyncio.to_thread(self.read, file, name=name, content_id=content_id)
        except Exception as e:
            log_error(f"Error reading file asynchronously: {str(e)}")
            return []

    def _extract_with_images(self, docx_document: Any, *, content_id: Optional[str]) -> str:
        if not content_id:
            raise ValueError("content_id is required when preserve_images=True")
        lines: List[str] = []

        def save_image(blob: bytes, media_type: str = "image/png") -> str:
            return save_image_markdown(
                content_id=content_id,
                data=blob,
                media_type=media_type,
                image_base_url=self.image_base_url,
            )

        for block in self._iter_block_items(docx_document):
            if isinstance(block, Paragraph):
                line = self._paragraph_with_inline_images(docx_document, block, save_image)
                if line:
                    lines.append(line)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            line = self._paragraph_with_inline_images(docx_document, para, save_image)
                            if line:
                                lines.append(line)

        return "\n\n".join(lines)

    @staticmethod
    def _iter_block_items(document: Any) -> Iterator[Union[Paragraph, Table]]:
        """Yield body paragraphs and tables in document order."""
        body = document.element.body
        for child in body.iterchildren():
            if child.tag == _W_P:
                yield Paragraph(child, document)
            elif child.tag == _W_TBL:
                yield Table(child, document)

    def _paragraph_with_inline_images(
        self,
        doc: Any,
        para: Any,
        save_image: Callable[[bytes, str], str],
    ) -> str:
        """Join run text with inline ``![](...)`` image links; not full markdown."""
        parts: List[str] = []
        for run in para.runs:
            parts.extend(self._run_to_parts(doc, run, save_image))
        return "".join(parts).strip()

    def _run_to_parts(
        self,
        doc: Any,
        run: Any,
        save_image: Callable[[bytes, str], str],
    ) -> List[str]:
        """Walk run XML in order so text and images stay interleaved."""
        parts: List[str] = []
        for child in run._element.iterchildren():
            if child.tag == _W_T:
                if child.text:
                    parts.append(child.text)
            elif child.tag == _W_TAB:
                parts.append("\t")
            elif child.tag in {_W_BR, _W_CR}:
                parts.append("\n")
            elif child.tag in {_W_DRAWING, _W_PICT}:
                for blob, media_type in self._extract_images_from_element(doc, child):
                    parts.append(save_image(blob, media_type))
            else:
                # Nested containers (e.g. w:r inside smart tags) — look for drawings/picts.
                for blob, media_type in self._extract_images_from_element(doc, child):
                    parts.append(save_image(blob, media_type))
        return parts

    def _extract_images_from_element(self, doc: Any, element: Any) -> List[Tuple[bytes, str]]:
        """Extract all resolvable images under a drawing/pict (or nested) element."""
        images: List[Tuple[bytes, str]] = []
        try:
            # DrawingML pictures (modern Word)
            for blip in element.xpath(".//a:blip"):
                resolved = self._resolve_image_part(
                    doc,
                    blip.get(_REL_EMBED) or blip.get(_REL_LINK),
                )
                if resolved is not None:
                    images.append(resolved)
        except Exception as e:
            log_error(f"Failed to extract DrawingML DOCX image: {e}")

        try:
            # Legacy VML pictures (`v` is not in python-docx's default nsmap).
            for imagedata in element.xpath(".//*[local-name()='imagedata']"):
                resolved = self._resolve_image_part(
                    doc,
                    imagedata.get(_REL_ID) or imagedata.get(_REL_EMBED),
                )
                if resolved is not None:
                    images.append(resolved)
        except Exception as e:
            log_error(f"Failed to extract VML DOCX image: {e}")

        return images

    @staticmethod
    def _resolve_image_part(doc: Any, r_id: Optional[str]) -> Optional[Tuple[bytes, str]]:
        if not r_id:
            return None
        related_parts = doc.part.related_parts
        if r_id not in related_parts:
            log_warning(f"DOCX image relationship not found: {r_id}")
            return None
        image_part = related_parts[r_id]
        blob = getattr(image_part, "blob", None)
        if blob is None:
            blob = getattr(image_part, "_blob", None)
        if not blob:
            return None
        content_type = getattr(image_part, "content_type", None) or "image/png"
        return blob, content_type
