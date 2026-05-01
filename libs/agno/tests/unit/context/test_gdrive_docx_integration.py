"""
Integration test for .docx text extraction in GDrive read_file.

Mocks the Google Drive API to verify the full extraction path works.
"""

from __future__ import annotations

import io
import json
from unittest.mock import MagicMock, patch

import pytest

from agno.context.gdrive.tools import DOCX_MIME_TYPE, AllDrivesGoogleDriveTools

try:
    import docx  # noqa: F401

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Bypass the @authenticate decorator for testing
def mock_authenticate(func):
    return func


@pytest.fixture
def docx_bytes() -> bytes:
    """Create a minimal valid .docx file for testing."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Hello from Google Drive!")
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def mock_drive_service(docx_bytes: bytes):
    """Mock Google Drive service that returns our test docx."""
    service = MagicMock()

    # Mock files().get() for metadata
    metadata = {
        "id": "test-file-id",
        "name": "test-document.docx",
        "mimeType": DOCX_MIME_TYPE,
        "size": str(len(docx_bytes)),
    }
    service.files().get().execute.return_value = metadata

    # Mock files().get_media() for content download
    media_request = MagicMock()
    media_request.execute.return_value = docx_bytes

    # MediaIoBaseDownload needs a request with an http attribute
    media_request.http = MagicMock()
    media_request.uri = "https://mock.googleapis.com/download"

    def mock_next_chunk():
        return (MagicMock(progress=lambda: 1.0), True)

    service.files().get_media.return_value = media_request

    return service, docx_bytes


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestDocxReadFileIntegration:
    """Test the full read_file path for .docx files."""

    def test_read_docx_extracts_text(self, mock_drive_service, docx_bytes: bytes):
        """Verify read_file extracts text from .docx when python-docx is available."""
        service, content = mock_drive_service

        tools = AllDrivesGoogleDriveTools()
        tools.service = service

        # Patch _download_bytes and bypass auth decorator
        with (
            patch("agno.context.gdrive.tools._download_bytes", return_value=content),
            patch("agno.context.gdrive.tools.authenticate", mock_authenticate),
        ):
            # Call the underlying method directly to bypass decorator
            result = AllDrivesGoogleDriveTools.read_file.__wrapped__(tools, "test-file-id")

        data = json.loads(result)

        assert "error" not in data
        assert "content" in data
        assert "Hello from Google Drive!" in data["content"]
        assert "test document" in data["content"]
        assert "multiple paragraphs" in data["content"]
        assert data["extractedFrom"] == "docx"
        assert data["file"]["name"] == "test-document.docx"

    def test_read_docx_respects_max_size(self, mock_drive_service):
        """Verify read_file rejects oversized .docx files."""
        service, content = mock_drive_service

        tools = AllDrivesGoogleDriveTools()
        tools.service = service
        tools.max_read_size = 100  # Very small limit

        # Update mock to return large file size
        service.files().get().execute.return_value = {
            "id": "test-file-id",
            "name": "huge-document.docx",
            "mimeType": DOCX_MIME_TYPE,
            "size": "1000000",  # 1MB
        }

        # Call the underlying method directly to bypass decorator
        result = AllDrivesGoogleDriveTools.read_file.__wrapped__(tools, "test-file-id")
        data = json.loads(result)

        assert "error" in data
        assert "exceeds max_read_size" in data["error"]


class TestDocxNotInstalled:
    """Test behavior when python-docx is not installed."""

    def test_shows_install_instructions_when_docx_missing(self):
        """Verify helpful error when python-docx is not available."""
        service = MagicMock()
        service.files().get().execute.return_value = {
            "id": "test-file-id",
            "name": "document.docx",
            "mimeType": DOCX_MIME_TYPE,
            "size": "1000",
        }

        tools = AllDrivesGoogleDriveTools()
        tools.service = service

        # Mock _download_bytes and make _extract_docx_text raise ImportError
        def raise_import_error(content_bytes):
            raise ImportError("No module named 'docx'")

        with (
            patch("agno.context.gdrive.tools._download_bytes", return_value=b"fake"),
            patch("agno.context.gdrive.tools._extract_docx_text", side_effect=raise_import_error),
        ):
            result = AllDrivesGoogleDriveTools.read_file.__wrapped__(tools, "test-file-id")

        data = json.loads(result)

        assert "error" in data
        assert "python-docx not installed" in data["error"]
        assert "pip install python-docx" in data["error"]
