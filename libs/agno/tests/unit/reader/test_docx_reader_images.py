"""Unit tests for DocxReader image preservation."""

from pathlib import Path

import pytest

from agno.knowledge.image import LocalKnowledgeImageStore, set_image_store
from agno.knowledge.reader.docx_reader import DocxReader

pil = pytest.importorskip("PIL.Image")


@pytest.fixture
def image_store(tmp_path: Path):
    store = LocalKnowledgeImageStore(base_dir=str(tmp_path / "doc_images"))
    set_image_store(store)
    yield store
    set_image_store(None)


def _png(tmp_path: Path, name: str = "figure.png") -> Path:
    image_path = tmp_path / name
    pil.new("RGB", (8, 8), color=(255, 0, 0)).save(image_path)
    return image_path


def _build_docx_with_inline_image(tmp_path: Path) -> Path:
    from docx import Document as DocxDocument

    image_path = _png(tmp_path)
    doc = DocxDocument()
    doc.add_paragraph("Before image")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(str(image_path))
    doc.add_paragraph("After image")

    out = tmp_path / "sample.docx"
    doc.save(out)
    return out


def _build_docx_with_table_image(tmp_path: Path) -> Path:
    from docx import Document as DocxDocument

    image_path = _png(tmp_path, "table.png")
    doc = DocxDocument()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Cell text"
    cell.paragraphs[0].add_run().add_picture(str(image_path))
    doc.add_paragraph("Outro")

    out = tmp_path / "table.docx"
    doc.save(out)
    return out


def _build_docx_with_text_and_image_same_run(tmp_path: Path) -> Path:
    from docx import Document as DocxDocument

    image_path = _png(tmp_path, "mixed.png")
    doc = DocxDocument()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Prefix ")
    run.add_picture(str(image_path))
    paragraph.add_run(" suffix")

    out = tmp_path / "mixed.docx"
    doc.save(out)
    return out


def test_docx_reader_preserve_images_inserts_inline_image_link(tmp_path: Path, image_store):
    docx_path = _build_docx_with_inline_image(tmp_path)
    reader = DocxReader(
        preserve_images=True,
        image_base_url="/knowledge/images",
        chunk=False,
    )

    documents = reader.read(docx_path, content_id="doc-1")
    assert len(documents) == 1
    content = documents[0].content
    assert "Before image" in content
    assert "After image" in content
    assert "![](/knowledge/images/doc-1/" in content
    saved = list((tmp_path / "doc_images" / "doc-1").glob("img-*"))
    assert saved


def test_docx_reader_preserve_images_in_table(tmp_path: Path, image_store):
    docx_path = _build_docx_with_table_image(tmp_path)
    reader = DocxReader(
        preserve_images=True,
        image_base_url="/knowledge/images",
        chunk=False,
    )

    content = reader.read(docx_path, content_id="doc-table")[0].content
    assert "Intro" in content
    assert "Cell text" in content
    assert "Outro" in content
    assert "![](/knowledge/images/doc-table/" in content
    assert list((tmp_path / "doc_images" / "doc-table").glob("img-*"))


def test_docx_reader_preserves_text_around_inline_image(tmp_path: Path, image_store):
    docx_path = _build_docx_with_text_and_image_same_run(tmp_path)
    reader = DocxReader(
        preserve_images=True,
        image_base_url="/knowledge/images",
        chunk=False,
    )

    content = reader.read(docx_path, content_id="doc-mixed")[0].content
    assert "Prefix" in content
    assert "suffix" in content
    assert "![](/knowledge/images/doc-mixed/" in content
    assert content.index("Prefix") < content.index("![")
    assert content.index("![") < content.index("suffix")


def test_docx_reader_preserve_images_default_off():
    reader = DocxReader()
    assert reader.preserve_images is False
